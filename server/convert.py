import sys
import re
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ============================================================
# Font handling
# ============================================================
LAO_RANGE = range(0x0E80, 0x0F00)

FONT_MAP = {
    'phetsarath': 'Phetsarath OT',
    'saysettha': 'Saysettha OT',
    'dokchampa': 'DokChampa',
    'laomn': 'Lao MN',
    'laosangam': 'Lao Sangam MN',
    'notosanslao': 'Noto Sans Lao',
    'notoseriflao': 'Noto Serif Lao',
    'arialmt': 'Arial',
    'arial': 'Arial',
    'timesnewromanpsmt': 'Times New Roman',
    'timesnewroman': 'Times New Roman',
    'calibri': 'Calibri',
    'wingdings': 'Wingdings',
    'courier': 'Courier New',
    'couriernew': 'Courier New',
    'helvetica': 'Arial',
}

DEFAULT_LAO_FONT = 'Phetsarath OT'
DEFAULT_FONT = 'Arial'


def contains_lao(text):
    return any(ord(c) in LAO_RANGE for c in text)


def resolve_font(pdf_font_name, text):
    if not pdf_font_name:
        return DEFAULT_LAO_FONT if contains_lao(text) else DEFAULT_FONT
    clean = re.sub(r'^[A-Z]{6}\+', '', pdf_font_name)
    key = clean.lower().replace(' ', '').replace('-', '').replace('_', '')
    for k, v in FONT_MAP.items():
        if k in key:
            return v
    if contains_lao(text):
        return DEFAULT_LAO_FONT
    base = re.sub(r'[-,]?(Bold|Italic|BoldItalic|Light|Regular|Medium|SemiBold|Heavy|Black|Thin)$',
                  '', clean, flags=re.IGNORECASE)
    return base.strip() if base.strip() else DEFAULT_FONT


def is_bold(font_name, flags):
    if flags & (1 << 18):
        return True
    return bool(font_name and any(w in font_name.lower() for w in ['bold', 'heavy', 'black']))


def is_italic(font_name, flags):
    if flags & (1 << 1):
        return True
    return bool(font_name and any(w in font_name.lower() for w in ['italic', 'oblique']))


# ============================================================
# OOXML helpers
# ============================================================
def set_run_font(run, font_name):
    run.font.name = font_name
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), font_name)


def make_run(paragraph, text, font_name, font_size, bold, italic, color):
    run = paragraph.add_run(text)
    resolved = resolve_font(font_name, text)
    set_run_font(run, resolved)
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rPr.append(OxmlElement('w:bCs'))
    if italic:
        run.italic = True
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rPr.append(OxmlElement('w:iCs'))
    if color and color != 0:
        r_v = (color >> 16) & 0xFF
        g_v = (color >> 8) & 0xFF
        b_v = color & 0xFF
        run.font.color.rgb = RGBColor(r_v, g_v, b_v)
    return run


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            elem = OxmlElement(f'w:{edge}')
            for a, v in kwargs[edge].items():
                elem.set(qn(f'w:{a}'), str(v))
            old = tcBorders.find(qn(f'w:{edge}'))
            if old is not None:
                tcBorders.remove(old)
            tcBorders.append(elem)


def set_table_borders(table, sz=6, color="000000"):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(borders)


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    old = tcPr.find(qn('w:shd'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_valign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    old = tcPr.find(qn('w:vAlign'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(va)


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_twips)))
    tcW.set(qn('w:type'), 'dxa')
    old = tcPr.find(qn('w:tcW'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcW)


def set_cell_margin(cell, top=20, bottom=20, left=50, right=50):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(mar)


# ============================================================
# Detect real table columns from border drawings
# ============================================================
def detect_real_columns(page, table_bbox):
    """Detect actual column X positions from vertical border lines in drawings."""
    drawings = page.get_drawings()
    tb_x0, tb_y0, tb_x1, tb_y1 = table_bbox

    vertical_lines = []
    for d in drawings:
        fill = d.get('fill')
        rect = d.get('rect')
        if not rect or not fill:
            continue
        # Black thin vertical rectangle = border
        w = rect.x1 - rect.x0
        h = rect.y1 - rect.y0
        if w < 2 and h > 15:  # vertical line
            x_mid = (rect.x0 + rect.x1) / 2
            if tb_x0 - 2 <= x_mid <= tb_x1 + 2:
                vertical_lines.append(x_mid)

    if not vertical_lines:
        return None

    # Cluster nearby X positions (within 3pt)
    vertical_lines.sort()
    clusters = []
    for x in vertical_lines:
        if clusters and abs(x - clusters[-1]) < 3:
            clusters[-1] = (clusters[-1] + x) / 2  # average
        else:
            clusters.append(x)

    return clusters


def detect_real_rows(page, table_bbox):
    """Detect actual row Y positions from horizontal border lines."""
    drawings = page.get_drawings()
    tb_x0, tb_y0, tb_x1, tb_y1 = table_bbox

    horiz_lines = []
    for d in drawings:
        fill = d.get('fill')
        rect = d.get('rect')
        if not rect or not fill:
            continue
        w = rect.x1 - rect.x0
        h = rect.y1 - rect.y0
        if h < 2 and w > 30:  # horizontal line spanning significant width
            y_mid = (rect.y0 + rect.y1) / 2
            if tb_y0 - 2 <= y_mid <= tb_y1 + 2:
                # Check if it spans most of the table width
                if rect.x1 - rect.x0 > (tb_x1 - tb_x0) * 0.3:
                    horiz_lines.append(y_mid)

    if not horiz_lines:
        return None

    horiz_lines.sort()
    clusters = []
    for y in horiz_lines:
        if clusters and abs(y - clusters[-1]) < 3:
            clusters[-1] = (clusters[-1] + y) / 2
        else:
            clusters.append(y)

    return clusters


def detect_header_bg(page, table_bbox, header_y0, header_y1):
    """Detect header row background color."""
    drawings = page.get_drawings()
    for d in drawings:
        fill = d.get('fill')
        rect = d.get('rect')
        if not fill or not rect or len(fill) < 3:
            continue
        rv = int(fill[0] * 255)
        gv = int(fill[1] * 255)
        bv = int(fill[2] * 255)
        if rv >= 250 and gv >= 250 and bv >= 250:
            continue
        if rv == 0 and gv == 0 and bv == 0:
            continue
        # Check if this colored rect is in the header area
        if rect.y0 >= header_y0 - 2 and rect.y1 <= header_y1 + 2:
            return f"{rv:02X}{gv:02X}{bv:02X}"
    return None


# ============================================================
# Get spans within a rectangle
# ============================================================
def get_spans_in_rect(all_blocks, x0, y0, x1, y1, margin=2):
    result = []
    for block in all_blocks:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                sb = span["bbox"]
                cx = (sb[0] + sb[2]) / 2
                cy = (sb[1] + sb[3]) / 2
                if x0 - margin <= cx <= x1 + margin and y0 - margin <= cy <= y1 + margin:
                    result.append(span)
    return result


def spans_to_lines(spans):
    """Group spans into visual lines by Y coordinate."""
    if not spans:
        return []
    spans = sorted(spans, key=lambda s: (s['bbox'][1], s['bbox'][0]))
    lines = []
    cur = [spans[0]]
    cur_y = (spans[0]['bbox'][1] + spans[0]['bbox'][3]) / 2
    for sp in spans[1:]:
        sp_y = (sp['bbox'][1] + sp['bbox'][3]) / 2
        if abs(sp_y - cur_y) < 4:
            cur.append(sp)
        else:
            lines.append(sorted(cur, key=lambda s: s['bbox'][0]))
            cur = [sp]
            cur_y = sp_y
    lines.append(sorted(cur, key=lambda s: s['bbox'][0]))
    return lines


def write_spans_to_paragraph(p, spans, force_bold=False):
    """Write spans to a paragraph preserving original formatting."""
    for span in spans:
        text = span.get('text', '')
        if not text:
            continue
        fname = span.get('font', '')
        fsize = span.get('size', 10)
        fl = span.get('flags', 0)
        col = span.get('color', 0)
        b = is_bold(fname, fl) or force_bold
        it = is_italic(fname, fl)
        make_run(p, text, fname, fsize, b, it, col)


# ============================================================
# Main conversion
# ============================================================
def convert(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = DEFAULT_LAO_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), DEFAULT_LAO_FONT)

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        pw = page.rect.width
        ph = page.rect.height

        # Page setup
        if page_num == 0:
            sec = doc.sections[0]
        else:
            sec = doc.add_section()
        sec.page_width = Emu(int(pw / 72 * 914400))
        sec.page_height = Emu(int(ph / 72 * 914400))
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.65)
        sec.right_margin = Inches(0.55)

        page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        all_blocks = page_dict.get("blocks", [])

        # Detect tables using border analysis
        tables_found = page.find_tables()

        # Build real table structures from drawings
        real_tables = []
        for pdf_table in tables_found.tables:
            tb = pdf_table.bbox

            col_xs = detect_real_columns(page, tb)
            row_ys = detect_real_rows(page, tb)

            if col_xs and row_ys and len(col_xs) >= 2 and len(row_ys) >= 2:
                real_tables.append({
                    'bbox': tb,
                    'col_xs': col_xs,
                    'row_ys': row_ys,
                    'pdf_table': pdf_table
                })
            else:
                # Fallback: use find_tables result directly
                real_tables.append({
                    'bbox': tb,
                    'col_xs': None,
                    'row_ys': None,
                    'pdf_table': pdf_table
                })

        table_bboxes = [t['bbox'] for t in real_tables]

        def is_in_any_table(bbox):
            bx = (bbox[0] + bbox[2]) / 2
            by = (bbox[1] + bbox[3]) / 2
            for tb in table_bboxes:
                if tb[0] - 5 <= bx <= tb[2] + 5 and tb[1] - 5 <= by <= tb[3] + 5:
                    return True
            return False

        # Extract images reliably using get_images + get_image_info
        page_images = []
        img_info_list = page.get_image_info()
        img_list = page.get_images(full=True)

        # Build xref -> image data map
        xref_data = {}
        for img in img_list:
            xref = img[0]
            try:
                base = pdf.extract_image(xref)
                if base and base.get("image"):
                    xref_data[xref] = base
            except Exception:
                pass

        # Match image info (has bbox) with image data (has bytes)
        for info in img_info_list:
            bbox = info.get('bbox', (0, 0, 0, 0))
            xref = info.get('xref', 0)
            img_bytes = None

            # Try by xref first
            if xref in xref_data:
                img_bytes = xref_data[xref]['image']
            elif xref_data:
                # Fallback: use first available image data
                for k, v in xref_data.items():
                    img_bytes = v['image']
                    del xref_data[k]
                    break

            if img_bytes:
                page_images.append({
                    'bbox': bbox,
                    'data': img_bytes,
                    'y': bbox[1]
                })

        # Also check text dict image blocks as fallback
        for block in all_blocks:
            if block.get("type") == 1:
                bbox = block.get("bbox", (0, 0, 0, 0))
                img_data = block.get("image")
                # Check if we already have this image (by position)
                already = False
                for pi in page_images:
                    if abs(pi['bbox'][1] - bbox[1]) < 5 and abs(pi['bbox'][0] - bbox[0]) < 5:
                        already = True
                        break
                if not already and img_data:
                    page_images.append({
                        'bbox': bbox,
                        'data': img_data,
                        'y': bbox[1]
                    })

        # Collect elements
        elements = []
        for pi in page_images:
            elements.append(("image", pi['y'], pi))

        for block in all_blocks:
            bbox = block.get("bbox", (0, 0, 0, 0))
            if block.get("type") == 0:
                if not is_in_any_table(bbox):
                    elements.append(("text", bbox[1], block))

        for rt in real_tables:
            elements.append(("table", rt['bbox'][1], rt))

        elements.sort(key=lambda e: e[1])

        for etype, ey, edata in elements:

            # ---- IMAGE ----
            if etype == "image":
                img_info = edata
                try:
                    img_data = img_info['data']
                    bbox = img_info['bbox']
                    if not img_data:
                        continue

                    img_w_pt = bbox[2] - bbox[0]
                    img_h_pt = bbox[3] - bbox[1]
                    img_w_in = img_w_pt / 72.0
                    img_h_in = img_h_pt / 72.0
                    max_w = (pw / 72.0) - 1.2
                    img_cx = (bbox[0] + bbox[2]) / 2

                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

                    # Position based on where image sits on the page
                    if img_cx < pw * 0.35:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif img_cx > pw * 0.65:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Preserve exact size, cap to max page width
                    if img_w_in > max_w:
                        scale = max_w / img_w_in
                        img_w_in = max_w
                        img_h_in = img_h_in * scale

                    run = p.add_run()
                    run.add_picture(io.BytesIO(img_data),
                                    width=Inches(img_w_in),
                                    height=Inches(img_h_in))
                except Exception:
                    pass

            # ---- TEXT ----
            elif etype == "text":
                block = edata
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    full = ''.join(s.get('text', '') for s in spans)
                    if not full.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(0)
                        continue

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0.5)
                    p.paragraph_format.space_before = Pt(0.5)

                    first_x = spans[0]['bbox'][0]
                    last_x = spans[-1]['bbox'][2]
                    mid = (first_x + last_x) / 2
                    page_mid = pw / 2

                    if abs(mid - page_mid) < 25 and first_x > pw * 0.15:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif first_x > pw * 0.55:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        left_margin_pt = 0.65 * 72
                        indent_pt = max(0, first_x - left_margin_pt)
                        if indent_pt > 10:
                            p.paragraph_format.left_indent = Pt(indent_pt)

                    for span in spans:
                        text = span.get('text', '')
                        if not text:
                            continue
                        make_run(p, text, span.get('font', ''), span.get('size', 11),
                                 is_bold(span.get('font', ''), span.get('flags', 0)),
                                 is_italic(span.get('font', ''), span.get('flags', 0)),
                                 span.get('color', 0))

            # ---- TABLE ----
            elif etype == "table":
                rt = edata
                col_xs = rt['col_xs']
                row_ys = rt['row_ys']
                tb = rt['bbox']

                if col_xs and row_ys:
                    # Build table from real border positions
                    num_cols = len(col_xs) - 1
                    num_rows = len(row_ys) - 1

                    if num_cols < 1 or num_rows < 1:
                        continue

                    tbl = doc.add_table(rows=num_rows, cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(tbl, sz=6, color="000000")

                    # Set column widths
                    for c in range(num_cols):
                        w_pt = col_xs[c + 1] - col_xs[c]
                        w_twips = int(w_pt / 72 * 1440)
                        tbl.columns[c].width = Emu(int(w_pt / 72 * 914400))

                    # Detect header bg color
                    header_bg = detect_header_bg(page, tb, row_ys[0], row_ys[1]) if num_rows > 0 else None

                    # Fill each cell
                    for r in range(num_rows):
                        for c in range(num_cols):
                            cell = tbl.cell(r, c)

                            # Cell rect in PDF coordinates
                            cx0 = col_xs[c]
                            cy0 = row_ys[r]
                            cx1 = col_xs[c + 1]
                            cy1 = row_ys[r + 1]

                            # Get original spans for this cell
                            cell_spans = get_spans_in_rect(all_blocks, cx0, cy0, cx1, cy1, margin=3)

                            # Clear default
                            for pp in cell.paragraphs:
                                pp.clear()

                            if cell_spans:
                                span_lines = spans_to_lines(cell_spans)
                                for li, sline in enumerate(span_lines):
                                    if li == 0:
                                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                                    else:
                                        p = cell.add_paragraph()

                                    p.paragraph_format.space_before = Pt(0.5)
                                    p.paragraph_format.space_after = Pt(0.5)

                                    # Detect alignment
                                    line_text = ''.join(s.get('text', '') for s in sline).strip()

                                    # Numbers -> right align
                                    if line_text and re.match(r'^[\d,.\s]+(?:ກີບ|kip)?$', line_text, re.IGNORECASE):
                                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    elif r == 0:  # Header row -> center
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        # Check if text is centered in cell
                                        txt_x0 = min(s['bbox'][0] for s in sline)
                                        txt_x1 = max(s['bbox'][2] for s in sline)
                                        txt_mid = (txt_x0 + txt_x1) / 2
                                        cell_mid = (cx0 + cx1) / 2
                                        cell_w = cx1 - cx0
                                        if cell_w > 40 and abs(txt_mid - cell_mid) < 5:
                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        elif txt_x0 > cx0 + cell_w * 0.5:
                                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                                    # Detect indent within cell
                                    if r > 0 and len(sline) > 0:
                                        span_x0 = sline[0]['bbox'][0]
                                        indent = span_x0 - cx0
                                        if indent > 15:
                                            p.paragraph_format.left_indent = Pt(indent - 5)

                                    write_spans_to_paragraph(p, sline, force_bold=(r == 0))
                            else:
                                if not cell.paragraphs:
                                    cell.add_paragraph()

                            # Cell formatting
                            border_s = {"sz": "6", "val": "single", "color": "000000", "space": "0"}
                            set_cell_border(cell, top=border_s, bottom=border_s, left=border_s, right=border_s)
                            set_cell_valign(cell, 'center')
                            set_cell_margin(cell, top=15, bottom=15, left=40, right=40)

                            # Header background
                            if r == 0 and header_bg:
                                set_cell_shading(cell, header_bg)

                            # Column width
                            w_pt = col_xs[c + 1] - col_xs[c]
                            set_cell_width(cell, int(w_pt / 72 * 1440))

                else:
                    # Fallback: use find_tables data
                    pdf_table = rt['pdf_table']
                    table_data = pdf_table.extract()
                    if not table_data or not table_data[0]:
                        continue

                    num_rows = len(table_data)
                    num_cols = len(table_data[0])

                    tbl = doc.add_table(rows=num_rows, cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(tbl, sz=6, color="000000")

                    for r in range(num_rows):
                        for c in range(num_cols):
                            cell = tbl.cell(r, c)
                            text = table_data[r][c]
                            if cell.paragraphs:
                                p = cell.paragraphs[0]
                            else:
                                p = cell.add_paragraph()
                            if text:
                                fn = 'SaysetthaOT' if contains_lao(text) else 'ArialMT'
                                make_run(p, text, fn, 10, r == 0, False, 0)

                            border_s = {"sz": "6", "val": "single", "color": "000000", "space": "0"}
                            set_cell_border(cell, top=border_s, bottom=border_s, left=border_s, right=border_s)

        # Page break
        if page_num < len(pdf) - 1:
            doc.add_page_break()

    doc.save(docx_path)
    pdf.close()
    print('OK')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python3 convert.py <input.pdf> <output.docx>', file=sys.stderr)
        sys.exit(1)
    try:
        convert(sys.argv[1], sys.argv[2])
    except Exception as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)
